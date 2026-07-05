"""Resume text extraction (PDF / DOCX / TXT).

v0.1 keeps this dead-simple: pull raw text out of the file with `pypdf` or
`python-docx`, normalise whitespace, and hand it off to the LLM extractor
(see :mod:`app.services.resume_service`). No layout heuristics, no OCR.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_MIME_EXTENSIONS: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


class DocumentParseError(Exception):
    """Raised when a document can't be parsed."""


@dataclass
class ResumeText:
    text: str
    format: str            # pdf / docx / txt
    page_count: int | None  # None for docx/txt


def _detect_format(filename: str | None, content_type: str | None) -> str:
    if content_type and content_type in SUPPORTED_MIME_EXTENSIONS:
        return SUPPORTED_MIME_EXTENSIONS[content_type]
    if filename:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext in {"pdf", "docx", "txt"}:
            return ext
    raise DocumentParseError(
        f"unsupported_format: filename={filename!r} content_type={content_type!r}"
    )


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise DocumentParseError(f"pdf_read_failed: {e}") from e

    if reader.is_encrypted:
        # pypdf allows an empty-password decrypt attempt for lightly-locked PDFs.
        try:
            reader.decrypt("")
        except Exception as e:  # noqa: BLE001
            raise DocumentParseError("pdf_encrypted") from e

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    return "\n\n".join(pages), len(reader.pages)


def _extract_docx(data: bytes) -> str:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise DocumentParseError(f"docx_read_failed: {e}") from e

    parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]

    # Include table content — resumes often use tables for layout.
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n".join(parts)


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("txt_decode_failed")


_WHITESPACE_RE = re.compile(r"[ \t]+")
_MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


def _normalise(text: str) -> str:
    text = _WHITESPACE_RE.sub(" ", text)
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


def extract_resume_text(
    *,
    data: bytes,
    filename: str | None,
    content_type: str | None,
) -> ResumeText:
    """Extract plain text from a resume file.

    Args:
        data: Raw file bytes.
        filename: Original filename (used for extension sniffing).
        content_type: MIME type (from the multipart part).

    Raises:
        DocumentParseError: On any unsupported format or extraction failure.
    """
    if not data:
        raise DocumentParseError("empty_file")

    fmt = _detect_format(filename, content_type)
    page_count: int | None = None

    if fmt == "pdf":
        text, page_count = _extract_pdf(data)
    elif fmt == "docx":
        text = _extract_docx(data)
    else:  # txt
        text = _extract_txt(data)

    text = _normalise(text)
    if not text:
        raise DocumentParseError("empty_extraction")

    return ResumeText(text=text, format=fmt, page_count=page_count)
